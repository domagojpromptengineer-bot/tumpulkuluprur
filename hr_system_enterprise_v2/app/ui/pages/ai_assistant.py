import streamlit as st
import google.generativeai as genai
import pandas as pd
import io
from datetime import date, timedelta, datetime
from io import BytesIO
try:
    from docx import Document
except ImportError:
    pass

from ...common.utils import get_conn, query_df, log_action, get_position_rank

# --- HELPER FUNCTIONS ---
def get_ai_model():
    api_key = st.session_state.get("api_key")
    if not api_key:
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-2.5-flash')

def fill_contract_template(docx_file, worker_data):
    doc = Document(docx_file)
    def replace_text(paragraph):
        for key, value in worker_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    for p in doc.paragraphs: replace_text(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: replace_text(p)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def parse_ai_schedule_to_df(ai_text):
    try:
        lines = ai_text.split('\n')
        table_lines = [line for line in lines if '|' in line]
        if not table_lines: return None
        table_str = '\n'.join(table_lines)
        df = pd.read_csv(io.StringIO(table_str), sep="|", engine='python')
        df.columns = [c.strip() for c in df.columns]
        df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
        df = df[~df.iloc[:, 0].astype(str).str.contains('---')]
        return df
    except Exception as e:
        print(f"Parsing error: {e}")
        return None

# --- MAIN RENDER ---
def render(role):
    st.header("🤖 AI HR Asistent")
    
    if not st.session_state.get("api_key"):
        st.warning("⚠️ Molimo unesite Google AI API ključ u bočnu traku.")

    tabs = st.tabs(["📅 AI Scheduler (Raspored)", "⚖️ Legal & Contract Agent (RAG)"])

    # ---------------------------------------------------------
    # TAB 1: AI SCHEDULER
    # ---------------------------------------------------------
    with tabs[0]:
        st.subheader("Generiranje i Uređivanje Rasporeda")
        conn = get_conn()
        
        sektori_df = query_df("SELECT id, naziv FROM sektor")
        sektori_options = {r['naziv']: r['id'] for _, r in sektori_df.iterrows()}
        sektor_naziv = st.selectbox("Odaberi Sektor", options=sektori_options.keys())
        sektor_id = sektori_options.get(sektor_naziv)
        
        start_date = st.date_input("Početak tjedna (Ponedjeljak)", value=date.today())
        end_date = start_date + timedelta(days=6)
        additional_constraints = st.text_area("Dodatne upute za AI")
        
        if st.button("Generiraj Prijedlog Rasporeda"):
            if not st.session_state.get("api_key"):
                st.error("Nedostaje API ključ.")
            elif not sektor_id:
                st.error("Odaberite sektor.")
            else:
                with st.spinner("AI analizira..."):
                    # 1. Dohvati Smjene (Format)
                    smjene_df = query_df(f"SELECT substr(pocetak, 1, 5) as p, substr(kraj, 1, 5) as k FROM smjene WHERE sektor_id={sektor_id}")
                    smjene_list = ", ".join([f"{r['p']}-{r['k']}" for _, r in smjene_df.iterrows()])
                    
                    # 2. Dohvati Radnike (Hijerarhija)
                    employees_df = query_df(f"""
                        SELECT r.id, r.ime, r.prezime, p.naziv_pozicije 
                        FROM radnici r LEFT JOIN pozicije p ON r.pozicija_id = p.id
                        WHERE r.sektor_id = {sektor_id} AND r.status_zaposlenja = 'aktivan'
                    """)
                    employees_df['rank'] = employees_df['naziv_pozicije'].apply(get_position_rank)
                    employees_df = employees_df.sort_values(by=['rank', 'prezime'])
                    employees_list = "\n".join([f"- {row['ime']} {row['prezime']} ({row['naziv_pozicije'] or 'N/A'})" for _, row in employees_df.iterrows()])
                    
                    # 3. Eventi
                    events_df = query_df(f"SELECT naziv, tip_eventa, pocetak, kraj, potrebno_osoblje FROM events WHERE date(pocetak) <= '{end_date}' AND date(kraj) >= '{start_date}'")
                    events_list = "\n".join([f"EVENT: {e['naziv']} ({e['pocetak']}). Staffing: {e['potrebno_osoblje']}" for _, e in events_df.iterrows()]) if not events_df.empty else "Nema evenata."

                    # 4. Prompt
                    final_prompt = f"""
                    Ti si planer hotela. Kreiraj raspored za sektor '{sektor_naziv}' za tjedan od {start_date}.
                    
                    DOSTUPNE SMJENE (Koristi ISKLJUČIVO ove formate): {smjene_list}, ili 'SLOBODAN'.
                    
                    RADNICI (Poredani hijerarhijski):
                    {employees_list}
                    
                    EVENTI:
                    {events_list}
                    
                    UPUTE: {additional_constraints}
                    
                    FORMAT IZLAZA:
                    Markdown tablica.
                    Prvi stupac: 'Zaposlenik' (Format: "Ime Prezime (Pozicija)").
                    Ostali stupci: Datumi u formatu 'YYYY-MM-DD' (npr. 2024-05-01).
                    Ćelije: Samo vrijeme smjene (npr. 07:00-15:00) ili SLOBODAN.
                    """
                    
                    try:
                        model = get_ai_model()
                        response = model.generate_content(final_prompt)
                        st.session_state['ai_schedule_raw'] = response.text
                        st.session_state['ai_schedule_sector_id'] = sektor_id
                    except Exception as e:
                        st.error(f"AI Greška: {e}")

        if 'ai_schedule_raw' in st.session_state:
            st.markdown("### 📋 AI Prijedlog")
            with st.expander("Vidi sirovi tekst"): st.text(st.session_state['ai_schedule_raw'])
            
            df_parsed = parse_ai_schedule_to_df(st.session_state['ai_schedule_raw'])
            
            if df_parsed is not None and not df_parsed.empty:
                edited_df = st.data_editor(df_parsed, use_container_width=True)
                
                if st.button("💾 Spremi Raspored u Bazu", type="primary"):
                    try:
                        radnici_map = query_df(f"SELECT id, ime, prezime FROM radnici WHERE sektor_id={st.session_state['ai_schedule_sector_id']}")
                        saved_count = 0
                        cur = conn.cursor()
                        date_cols = edited_df.columns[1:]
                        
                        for index, row in edited_df.iterrows():
                            raw_name = row[edited_df.columns[0]]
                            clean_name = raw_name.split('(')[0].strip()
                            
                            radnik_id = None
                            for _, r in radnici_map.iterrows():
                                db_name = f"{r['ime']} {r['prezime']}"
                                if clean_name.lower() in db_name.lower() or db_name.lower() in clean_name.lower():
                                    radnik_id = r['id']
                                    break
                            
                            if radnik_id:
                                for col_date in date_cols:
                                    shift_val = row[col_date]
                                    try:
                                        target_date = col_date.strip()
                                        datetime.strptime(target_date, "%Y-%m-%d")
                                        
                                        # ROBUSNA PROVJERA VRIJEDNOSTI
                                        clean_val = None
                                        if shift_val and pd.notna(shift_val):
                                            s_val = str(shift_val).strip()
                                            if s_val.lower() not in ['none', 'nan', '', 'slobodan']:
                                                clean_val = s_val
                                        
                                        cur.execute("SELECT id FROM rasporedi WHERE radnik_id=? AND datum=?", (radnik_id, target_date))
                                        exists = cur.fetchone()
                                        
                                        if exists:
                                            if clean_val:
                                                cur.execute("UPDATE rasporedi SET opis_smjene=?, ai_recommended=1 WHERE id=?", (clean_val, exists['id']))
                                            else:
                                                cur.execute("DELETE FROM rasporedi WHERE id=?", (exists['id'],))
                                        elif clean_val:
                                            cur.execute("INSERT INTO rasporedi (sektor_id, radnik_id, datum, opis_smjene, ai_recommended) VALUES (?,?,?,?,1)",
                                                      (st.session_state['ai_schedule_sector_id'], radnik_id, target_date, clean_val))
                                        saved_count += 1
                                    except ValueError: pass
                        conn.commit()
                        st.success(f"Spremljeno {saved_count} zapisa!")
                    except Exception as e:
                        st.error(f"Greška: {e}")
            else:
                st.warning("Greška u parsiranju tablice.")
        conn.close()

    # ---------------------------------------------------------
    # TAB 2: LEGAL AGENT
    # ---------------------------------------------------------
    with tabs[1]:
        st.subheader("⚖️ Pravni Asistent & Generator Ugovora")
        mode = st.radio("Odaberite način rada:", ["💬 Pravni Chat (ZOR)", "📄 Generator Ugovora (Auto-Fill)"])
        
        if mode == "💬 Pravni Chat (ZOR)":
            user_q = st.text_input("Vaše pitanje:")
            if user_q and st.button("Pitaj AI"):
                if not st.session_state.get("api_key"): st.error("Nedostaje API ključ.")
                else:
                    try:
                        model = get_ai_model()
                        res = model.generate_content(f"Ti si pravni stručnjak za ZOR RH. Pitanje: {user_q}")
                        st.markdown(res.text)
                    except Exception as e: st.error(f"Greška: {e}")

        elif mode == "📄 Generator Ugovora (Auto-Fill)":
            uploaded_file = st.file_uploader("Uploadaj predložak (.docx)", type=["docx"])
            conn = get_conn()
            radnici_df = query_df("SELECT id, ime, prezime, oib, adresa, email FROM radnici")
            radnici_options = {f"{r['ime']} {r['prezime']}": r for _, r in radnici_df.iterrows()}
            sel_radnik = st.selectbox("Odaberi radnika", options=radnici_options.keys())
            
            if uploaded_file and sel_radnik and st.button("Generiraj"):
                r_data = radnici_options[sel_radnik]
                u_info = query_df(f"SELECT bruto FROM ugovori WHERE radnik_id={r_data['id']} ORDER BY id DESC LIMIT 1")
                placa = str(u_info.iloc[0]['bruto']) if not u_info.empty else "N/A"
                
                data_map = {
                    "ime": r_data['ime'], "prezime": r_data['prezime'], "oib": r_data['oib'],
                    "adresa": r_data['adresa'], "placa": placa, "datum": date.today().strftime("%d.%m.%Y")
                }
                try:
                    doc = fill_contract_template(uploaded_file, data_map)
                    st.download_button("⬇️ Preuzmi", data=doc, file_name=f"Ugovor_{r_data['prezime']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e: st.error(f"Greška: {e}")
            conn.close()