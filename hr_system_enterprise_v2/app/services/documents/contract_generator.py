from datetime import date
from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

def get_contract_template(data: dict) -> str:
    """
    Generira kompletan tekst Ugovora o radu sukladno Zakonu o radu RH.
    """
    # Pomoćne varijable za logiku
    danas = date.today().strftime('%d.%m.%Y.')
    mjesto = data.get('company_address', 'Zagrebu').split(',')[0].strip()
    
    # Logika za određeno/neodređeno
    tip_ugovora = data.get("tip_ugovora", "na neodređeno vrijeme")
    clanak_trajanje = "1. Ovaj Ugovor sklapa se na neodređeno vrijeme."
    if "određeno" in tip_ugovora.lower():
        kraj = data.get("kraj", "do povratka radnika s bolovanja")
        clanak_trajanje = f"1. Ovaj Ugovor sklapa se na određeno vrijeme do {kraj}, zbog povećanog opsega posla."

    # Formatiranje iznosa
    try:
        bruto = float(data.get('bruto', 0))
        bruto_str = f"{bruto:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        bruto_str = str(data.get('bruto', "0,00"))

    # TEKST UGOVORA
    return f"""
UGOVOR O RADU
sklopljen u {mjesto}, dana {danas}

Između:

1. POSLODAVAC:
   {data.get('company_name', '[NAZIV TVRTKE]')}
   OIB: {data.get('company_oib', '[OIB]')}
   Adresa: {data.get('company_address', '[ADRESA]')}
   Zastupan po: {data.get('company_director', 'Upravi')}
   (u daljnjem tekstu: Poslodavac)

2. RADNIK:
   {data.get('emp_name', '[IME I PREZIME]')}
   OIB: {data.get('emp_oib', '[OIB RADNIKA]')}
   Adresa: {data.get('emp_address', '[ADRESA RADNIKA]')}
   (u daljnjem tekstu: Radnik)


Članak 1. (Predmet ugovora)
1. Ovim Ugovorom zasniva se radni odnos između Poslodavca i Radnika.
2. Radnik će obavljati poslove radnog mjesta: {data.get('emp_position', 'Radnik')}.
3. Opis poslova radnog mjesta sastavni je dio Pravilnika o organizaciji i sistematizaciji radnih mjesta Poslodavca, s kojim je Radnik upoznat prije potpisivanja ovog Ugovora.

Članak 2. (Mjesto rada)
1. Mjesto rada je sjedište Poslodavca na adresi: {data.get('company_address')}.
2. Poslodavac može, zbog potrebe posla, uputiti Radnika na rad izvan sjedišta ili na terenski rad, sukladno zakonskim propisima.

Članak 3. (Trajanje radnog odnosa)
{clanak_trajanje}
2. Radnik počinje s radom dana: {data.get('pocetak', '[DATUM POČETKA]')}.

Članak 4. (Probni rad)
1. Ugovara se probni rad u trajanju od 3 (tri) mjeseca.
2. Ako Radnik ne zadovolji na probnom radu, otkazni rok iznosi 7 dana.

Članak 5. (Radno vrijeme)
1. Radnik će raditi u punom radnom vremenu od 40 sati tjedno.
2. Tjedno radno vrijeme raspoređuje se u pet radnih dana, od ponedjeljka do petka, osim ako narav posla ne zahtijeva drugačiji raspored (smjenski rad, rad vikendom), o čemu Poslodavac donosi pisanu odluku.
3. Radnik ima pravo na stanku (pauzu) od 30 minuta dnevno, koja se ubraja u radno vrijeme.

Članak 6. (Plaća i naknade)
1. Za obavljeni rad Poslodavac će Radniku isplaćivati osnovnu bruto plaću u iznosu od {bruto_str} EUR mjesečno.
2. Plaća se isplaćuje najkasnije do 15. u mjesecu za prethodni mjesec na tekući račun Radnika.
3. Radnik ima pravo na uvećanje plaće za prekovremeni rad, rad noću, rad nedjeljom i blagdanom, sukladno Zakonu o radu i Pravilniku o radu.

Članak 7. (Godišnji odmor)
1. Radnik ima pravo na plaćeni godišnji odmor u trajanju od najmanje 4 (četiri) tjedna za svaku kalendarsku godinu.
2. Raspored korištenja godišnjeg odmora utvrđuje Poslodavac, uzimajući u obzir potrebe organizacije rada i mogućnosti za odmor i rekreaciju Radnika.

Članak 8. (Otkaz ugovora)
1. Ovaj Ugovor može prestati sporazumom stranaka, otkazom ili istekom vremena na koje je sklopljen.
2. Otkazni rokovi utvrđuju se sukladno odredbama Zakona o radu.
3. Otkaz mora biti u pisanom obliku.

Članak 9. (Zaštita na radu)
1. Poslodavac je dužan osigurati uvjete za rad na siguran način i zaštitu zdravlja Radnika.
2. Radnik je dužan pridržavati se propisanih mjera zaštite na radu i koristiti osobna zaštitna sredstva.

Članak 10. (Poslovna tajna i zaštita podataka)
1. Radnik se obvezuje čuvati kao poslovnu tajnu sve podatke koje sazna tijekom rada, a koji su određeni kao poslovna tajna ili čije bi otkrivanje moglo štetiti Poslodavcu.
2. Obveza čuvanja poslovne tajne traje i nakon prestanka radnog odnosa.

Članak 11. (Završne odredbe)
1. Na prava i obveze koje nisu uređene ovim Ugovorom primjenjuju se odredbe Zakona o radu, Kolektivnog ugovora i Pravilnika o radu.
2. Sve sporove proizašle iz ovog Ugovora stranke će nastojati riješiti mirnim putem, a u slučaju nemogućnosti dogovora, nadležan je sud u {mjesto}.
3. Ovaj Ugovor sastavljen je u 2 (dva) istovjetna primjerka, od kojih svaka strana zadržava po 1 (jedan).


POSLODAVAC:                                         RADNIK:
_________________________                           _________________________
(Potpis i pečat)                                    (Vlastoručni potpis)
"""

def generate_contract_docx(contract_text: str) -> BytesIO:
    """Generira DOCX datoteku iz teksta."""
    doc = Document()
    
    # Naslov
    title = doc.add_paragraph("UGOVOR O RADU")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    
    # Tijelo teksta
    for line in contract_text.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)
            if "Članak" in line or "POSLODAVAC:" in line:
                p.runs[0].bold = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer